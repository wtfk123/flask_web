import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import jieba.posseg as pseg
from docx import Document
from collections import Counter
import pymongo
import networkx as nx
import sys
from pylab import mpl


mpl.rcParams["font.sans-serif"] = ["SimHei"]
mpl.rcParams["axes.unicode_minus"] = False


# 建立与MongoDB的连接
client = pymongo.MongoClient("mongodb://localhost:27017")
# 选择要操作的数据库
db = client["flask"]
# 选择要操作的集合
collection = db["Follower Colloctor"]
# 将集合赋值给变量data
data = collection.find()

# 将查询结果转换为列表
data_list = list(data)
# 将列表转换为DataFrame对象
data = pd.json_normalize(data_list)


input_text = sys.argv[1]
# 取出指定用户
username = 'https://twitter.com/' + input_text
D_data = data[data['query'] == username]
D_data['tweetsCount']=D_data['tweetsCount'].astype(int)
top_15_user = D_data.nlargest(25, 'tweetsCount')

G = nx.DiGraph()
G_path = 'D:/vscode/.venv/web/static/social.png'

# 添加主点ablinken
G.add_node(input_text)

# 添加其他节点和权重
for index, row in top_15_user.iterrows():
    screen_name = row['screenName']
    tweets_count = row['tweetsCount']
    G.add_node(screen_name)
    G.add_edge(input_text, screen_name, weight=tweets_count)

# 绘制拓扑图
pos = nx.spring_layout(G)
labels = nx.get_edge_attributes(G, 'weight')
node_sizes = [1500 if node != input_text else 4000 for node in G.nodes]  # 设置节点大小

nx.draw_networkx(G, pos=pos, with_labels=True, node_color='lightblue', node_size=node_sizes, font_size=9)
nx.draw_networkx_edge_labels(G, pos=pos, edge_labels=labels)

plt.title("社交网络关系拓扑图")
plt.tight_layout()
plt.savefig(G_path)

# 创建空文档L
doc = Document()

# 插入图到 Word 文档
doc.add_paragraph('指定用户的社交网络拓扑图')
doc.add_picture(G_path)

# 保存文档
doc.save('./social.docx')